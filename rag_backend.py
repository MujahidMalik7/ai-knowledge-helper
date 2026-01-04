import os
import pandas as pd
import pytesseract
from io import BytesIO
from pypdf import PdfReader
from pdf2image import convert_from_path
from docx import Document
from sentence_transformers import SentenceTransformer, CrossEncoder
import faiss
import numpy as np
import ollama

model = SentenceTransformer("BAAI/bge-base-en-v1.5")
reranker = CrossEncoder('cross-encoder/ms-marco-MiniLM-L-6-v2')
# LLM
llm = "mistral:7b"

# chat_history = []

def extract_text(file_input):
    """
    Extract text from either a file path (str) or Streamlit UploadedFile object.
    Keeps your exact original logic, including OCR fallback for scanned PDFs.
    """
    # Determine if input is UploadedFile or file path
    if hasattr(file_input, "read"):  # It's a Streamlit UploadedFile
        file_name = file_input.name
        file_bytes = file_input.getvalue()  # Read all bytes
        file_input.seek(0)  # Reset pointer in case needed elsewhere
        is_uploaded = True
    else:
        # It's a regular file path (str)
        file_name = os.path.basename(file_input)
        file_bytes = None
        is_uploaded = False

    ext = os.path.splitext(file_name)[1].lower()
    text = ""

    # -------- PDF --------
    if ext == ".pdf":
        if is_uploaded:
            # Read from bytes
            reader = PdfReader(BytesIO(file_bytes))
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
        else:
            # Read from file path
            reader = PdfReader(file_input)
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"

        # OCR fallback if no text extracted (likely scanned PDF)
        if text.strip() == "":
            if is_uploaded:
                # Convert PDF bytes to images
                from pdf2image import convert_from_bytes
                images = convert_from_bytes(file_bytes)
            else:
                images = convert_from_path(file_input)

            for img in images:
                text += pytesseract.image_to_string(img) + "\n"

    # -------- DOCX --------
    elif ext == ".docx":
        if is_uploaded:
            doc = Document(BytesIO(file_bytes))
        else:
            doc = Document(file_input)
        text = "\n".join([p.text for p in doc.paragraphs])

    # -------- CSV --------
    elif ext == ".csv":
        if is_uploaded:
            df = pd.read_csv(BytesIO(file_bytes))
        else:
            df = pd.read_csv(file_input)
        text = df.to_string(index=False)

    # -------- TXT --------
    elif ext == ".txt":
        if is_uploaded:
            text = file_bytes.decode("utf-8", errors="ignore")
        else:
            with open(file_input, "r", encoding="utf-8", errors="ignore") as f:
                text = f.read()

    else:
        raise ValueError(f"Unsupported file type: {ext}")

    return text.strip()

def chunk_text(text, chunk_size=200, overlap=40):
    words = text.split()
    chunks = []

    start = 0
    while start < len(words):
        end = start + chunk_size
        chunk = words[start:end]
        chunks.append(" ".join(chunk))
        start += chunk_size - overlap

    return chunks

def embed_chunks(chunks):
    """
    Input: list of text chunks
    Output: numpy array of embeddings
    """
    embeddings = model.encode(chunks, convert_to_numpy=True, show_progress_bar=True)
    return embeddings

def build_faiss_index(embeddings):
    """
    Input: numpy array of embeddings
    Output: FAISS index
    """
    dim = embeddings.shape[1]   # embedding size of BGE
    index = faiss.IndexFlatIP(dim)
    
    # Normalize embeddings for cosine similarity
    faiss.normalize_L2(embeddings)
    
    index.add(embeddings)
    print(f"FAISS index built with {index.ntotal} vectors")
    return index

def retrieve_and_rerank(query, index, chunks, top_k_faiss=10, top_n=5):
    """
    query      : str, user query
    index      : FAISS index
    chunks     : list of all text chunks
    top_k_faiss: int, how many FAISS results to fetch
    top_n     : int, how many final results to return after reranking
    """

    # Embed query
    query_embedding = model.encode([query], convert_to_numpy=True)
    faiss.normalize_L2(query_embedding)

    # Retrieve top_k from FAISS
    D, I = index.search(query_embedding, top_k_faiss)  # I = indices, D = scores
    retrieved_chunks = [chunks[i] for i in I[0]]

    # Prepare pairs for CrossEncoder (query, chunk)
    pairs = [[query, chunk] for chunk in retrieved_chunks]

    # Rerank using CrossEncoder
    scores = reranker.predict(pairs)
    
    # Sort by score descending
    ranked_chunks = sorted(zip(retrieved_chunks, scores), key=lambda x: x[1], reverse=True)

    # Return top_n results
    top_chunks = [chunk for chunk, score in ranked_chunks[:top_n]]
    top_scores = [score for chunk, score in ranked_chunks[:top_n]]

    return top_chunks, top_scores

def generate_answer_ollama(query, top_chunks, conversation_history=None):
    """
    Generate answer using Ollama with RAG context
    
    Args:
        query: User's question
        top_chunks: Retrieved document chunks
        conversation_history: Optional list of dicts with 'query' and 'answer' keys
    """
    if not top_chunks:
        context = "No relevant information found in the document."
    else:
        context = "\n".join(top_chunks)

    # Build messages for Ollama
    messages = []
    
    # System message
    messages.append({
        "role": "system",
        "content": (
            "You are a helpful assistant that answers questions based on the provided document context. "
            "Be concise, accurate, and friendly. "
            "If the question refers to previous messages, use conversation history. "
            "If no relevant info in context, say you couldn't find it in the document."
        )
    })
    
    # Add conversation history if provided
    if conversation_history:
        for item in conversation_history:
            messages.append({"role": "user", "content": item['query']})
            messages.append({"role": "assistant", "content": item['answer']})
    
    # Add current query with context
    user_message = f"""Context from document:
{context}

Question: {query}

Please answer based on the context provided."""
    
    messages.append({
        "role": "user",
        "content": user_message
    })

    try:
        response = ollama.chat(
            model=llm,
            messages=messages
        )
        answer = response['message']['content'].strip()
        return answer

    except Exception as e:
        return f"Error generating answer: {e}"
    
    
def download_chat_history(chat_history, filename="chat_history.txt"):
    """
    Creates a downloadable text file from chat history.
    
    Args:
        chat_history: List of dicts with 'query' and 'answer' keys
        filename: Name of the output file
    
    Returns:
        String containing formatted chat history
    """
    if not chat_history:
        return "No chat history available."
    
    content = "=" * 60 + "\n"
    content += "CHAT HISTORY\n"
    content += "=" * 60 + "\n\n"
    
    for i, item in enumerate(chat_history, 1):
        content += f"Q{i}: {item['query']}\n"
        content += f"A{i}: {item['answer']}\n"
        content += "-" * 60 + "\n\n"    
    
    content += "=" * 60 + "\n"
    content += f"Total interactions: {len(chat_history)}\n"
    content += "=" * 60 + "\n"
    
    return content

def document_summary(text, max_length=500):
    """
    Generate a concise summary of the entire document using Ollama.
    
    Args:
        text: Full document text (string)
        max_length: Maximum number of words to include from the original text for summarization
    
    Returns:
        String containing the summary
    """
    if not text or not text.strip():
        return "No text available to summarize."
    
    # If text is too long, take a representative sample from beginning, middle, and end
    words = text.split()
    if len(words) > max_length:
        # Take first 40%, middle 30%, last 30% of max_length
        start_size = int(max_length * 0.4)
        middle_size = int(max_length * 0.3)
        end_size = max_length - start_size - middle_size
        
        middle_start = len(words) // 2 - middle_size // 2
        
        sample_text = (
            " ".join(words[:start_size]) + " [...] " +
            " ".join(words[middle_start:middle_start + middle_size]) + " [...] " +
            " ".join(words[-end_size:])
        )
    else:
        sample_text = text
    
    # Create prompt for summarization
    prompt = f"""Please provide a concise summary of the following document. 
The summary should be 3-5 sentences that capture the main topics, key points, and overall theme.

Document:
{sample_text}

Summary:"""
    
    try:
        response = ollama.chat(
            model=llm,
            messages=[
                {
                    "role": "system",
                    "content": "You are a helpful assistant that creates clear, concise summaries of documents. Focus on the main ideas and key information."
                },
                {
                    "role": "user",
                    "content": prompt
                }
            ]
        )
        summary = response['message']['content'].strip()
        return summary
    
    except Exception as e:
        return f"Error generating summary: {e}"